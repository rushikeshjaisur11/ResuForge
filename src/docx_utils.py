"""
CLI utility for reading and writing .docx files.
Called by Claude skills — no AI logic here, purely mechanical.

Usage:
  python src/docx_utils.py read <resume.docx>
      Prints JSON: {"full_text": "...", "paragraphs": [{"index": int, "text": str, "style": str}]}

  python src/docx_utils.py write <src.docx> <out.docx> <mapping.json>
      Applies text mapping to src.docx and saves to out.docx.
      mapping.json: {"original line": "rewritten line", ...}

  python src/docx_utils.py write-sections <src.docx> <out.docx> <sections.json>
      Applies section-level rewrites. Supports rearranging, adding, and removing bullets.
      sections.json format:
        {
          "sections": [
            {"type": "text_replace", "original": "old text", "replacement": "new text"},
            {
              "type": "role_bullets",
              "original_bullets": ["exact bullet text 1", "exact bullet text 2"],
              "bullets": ["new bullet 1", "new bullet 2", "new bullet 3"]
            }
          ]
        }
"""

import json
import os
import sys
from copy import deepcopy

from docx import Document


def cmd_read(path: str) -> None:
    doc = Document(path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": para.style.name,
            })
    full_text = "\n".join(p["text"] for p in paragraphs)
    print(json.dumps({"full_text": full_text, "paragraphs": paragraphs}, ensure_ascii=False, indent=2))


def cmd_write(src_path: str, out_path: str, mapping_path: str) -> None:
    with open(mapping_path, encoding="utf-8") as f:
        mapping = json.load(f)

    doc = Document(src_path)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    for para in doc.paragraphs:
        if para.text in mapping:
            _replace_para(para, mapping[para.text])
        else:
            for run in para.runs:
                if run.text and run.text in mapping:
                    run.text = mapping[run.text]

    doc.save(out_path)
    print(f"Saved: {out_path}")


def cmd_write_sections(src_path: str, out_path: str, sections_path: str) -> None:
    with open(sections_path, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document(src_path)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    paras = doc.paragraphs  # snapshot — elements stay valid after XML mutations

    for section in data.get("sections", []):
        stype = section.get("type")

        if stype == "text_replace":
            original = section["original"].strip()
            for para in paras:
                if para.text.strip() == original:
                    _replace_para(para, section["replacement"])
                    break

        elif stype == "role_bullets":
            _apply_role_bullets(paras, section)


    doc.save(out_path)
    print(f"Saved: {out_path}")


def _apply_role_bullets(paras: list, section: dict) -> None:
    """Replace, reorder, add, or remove bullets for one role.

    Finds each paragraph in `original_bullets` by exact text match (in document
    order), then replaces their content with `bullets`. Extra originals are
    deleted; extra new bullets are inserted after the last matched paragraph,
    using the first matched paragraph's formatting as a template.
    """
    original_texts = [b.strip() for b in section.get("original_bullets", [])]
    new_bullets = section["bullets"]

    if not original_texts:
        print("Warning: 'original_bullets' is empty — skipping role_bullets section", file=sys.stderr)
        return

    # Find original bullet paragraphs in document order, matching each text once
    remaining = list(original_texts)
    bullet_paras = []
    for para in paras:
        if para.text.strip() in remaining:
            bullet_paras.append(para)
            remaining.remove(para.text.strip())
        if not remaining:
            break

    if not bullet_paras:
        print(f"Warning: none of the original_bullets were found in the document", file=sys.stderr)
        return
    if remaining:
        print(f"Warning: some original_bullets not found (skipped): {remaining}", file=sys.stderr)

    # Deep-copy the first bullet's XML element as a formatting template for new inserts
    template_ele = deepcopy(bullet_paras[0]._element)

    n_existing = len(bullet_paras)
    n_new = len(new_bullets)

    # Replace text in existing paragraph slots
    for j in range(min(n_existing, n_new)):
        _replace_para(bullet_paras[j], new_bullets[j])

    if n_new > n_existing:
        # Insert additional bullets after the last existing one
        last_ele = bullet_paras[-1]._element
        for j in range(n_existing, n_new):
            new_ele = deepcopy(template_ele)
            _set_element_text(new_ele, new_bullets[j])
            last_ele.addnext(new_ele)
            last_ele = new_ele

    elif n_existing > n_new:
        # Remove extra original paragraphs
        for j in range(n_new, n_existing):
            ele = bullet_paras[j]._element
            ele.getparent().remove(ele)


def _set_element_text(ele, new_text: str) -> None:
    """Set text on a raw paragraph XML element (used for deepcopied template elements)."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    runs = ele.findall(f".//{{{W}}}r")
    if not runs:
        return
    t_eles = runs[0].findall(f"{{{W}}}t")
    if t_eles:
        t_eles[0].text = new_text
        t_eles[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    for run in runs[1:]:
        run.getparent().remove(run)


def _replace_para(para, new_text: str) -> None:
    """Put new_text in the first run, clear the rest (preserves formatting)."""
    if not para.runs:
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    command = sys.argv[1]
    if command == "read":
        cmd_read(sys.argv[2])
    elif command == "write":
        if len(sys.argv) < 5:
            print("Usage: python src/docx_utils.py write <src.docx> <out.docx> <mapping.json>")
            sys.exit(1)
        cmd_write(sys.argv[2], sys.argv[3], sys.argv[4])
    elif command == "write-sections":
        if len(sys.argv) < 5:
            print("Usage: python src/docx_utils.py write-sections <src.docx> <out.docx> <sections.json>")
            sys.exit(1)
        cmd_write_sections(sys.argv[2], sys.argv[3], sys.argv[4])
    else:
        print(f"Unknown command: {command}")
        sys.exit(1)
